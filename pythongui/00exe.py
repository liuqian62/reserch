# coding=gbk
import os
import re
import docx
import pandas as pd
import tkinter as tk
from docx        import Document
from tkinter     import filedialog
from win32com    import client as wc
from prettytable import PrettyTable
from openpyxl    import Workbook
from openpyxl    import load_workbook


##������ʾ
def play(x):
    winNew = tk.Toplevel(w)
    winNew.geometry('1400x800')
    winNew.title('���ļ�����')

    scro1 = tk.Scrollbar(winNew)
    scro1.pack(side=tk.RIGHT, fill=tk.Y)
    scro2 = tk.Scrollbar(winNew,orient=tk.HORIZONTAL)
    scro2.pack(side=tk.BOTTOM, fill=tk.X)
    text = tk.Text(winNew, width=700, height=800)
    text.pack(side=tk.LEFT, fill=tk.BOTH)
    scro1.config(command=text.yview)
    scro2.config(command=text.xview)
    text.config(xscrollcommand=scro2.set,yscrollcommand=scro1.set)
    text.pack(side=tk.LEFT,fill=tk.BOTH) 

    text.delete(1.0, "end")
    text.insert(tk.INSERT, x) 
 
    
#����Ʒ�����ѯ
def classify_product(data, product):
    x = PrettyTable(["Item",
                      "Comment",
                      "Discription",
                      "Designator",
                      "Footprint",
                      "LibRef",
                      "Quantity",
                      "Manufacturer",
                      "Note",
                      "Level",
                      "Product"])
    for index, row in data.iterrows():
        if row["Product"] == product:
            p_list = row.values.tolist()
            x.add_row(p_list)

    play(x)



#��������������ѯ
def classify_num(data, min, max):
    x = PrettyTable(["Item",
                      "Comment",
                      "Discription",
                      "Designator",
                      "Footprint",
                      "LibRef",
                      "Quantity",
                      "Manufacturer",
                      "Note",
                      "Level",
                      "Product"])
    for index, row in data.iterrows():
        if row["Quantity"] >= min and row["Quantity"] <= max:
            p_list = row.values.tolist()
            x.add_row(p_list)

    play(x)

#����������ѯ
def classify_type(data, type):
    x = PrettyTable(["Item",
                      "Comment",
                      "Discription",
                      "Designator",
                      "Footprint",
                      "LibRef",
                      "Quantity",
                      "Manufacturer",
                      "Note",
                      "Level",
                      "Product"])
    for index, row in data.iterrows():
        str = row["Designator"]
        t = ""
        for i in str:
            if i.isalpha() == True:
                t = t + i
            if i.isalpha() == False:
                break
        if t == type:
            p_list = row.values.tolist()
            x.add_row(p_list)

    play(x)


#�������������ѯ
def classify_multi(data, product, min, max, type):
    x = PrettyTable(["Item",
                      "Comment",
                      "Discription",
                      "Designator",
                      "Footprint",
                      "LibRef",
                      "Quantity",
                      "Manufacturer",
                      "Note",
                      "Level",
                      "Product"])
    for index, row in data.iterrows():
        str = row["Designator"]
        t = ""
        for i in str:
            if i.isalpha() == True:
                t = t + i
            if i.isalpha() == False:
                break
        if (row["Product"] == product
            and row["Quantity"] >= min and row["Quantity"] <= max
            and t == type):
            p_list = row.values.tolist()
            x.add_row(p_list)

    play(x)


    
##���ļ���ѯ����
def file():
    winNew = tk.Toplevel(w)
    winNew.geometry('700x350')
    winNew.title('���ļ�����')
    t1=tk.Entry(winNew,width=20,textvariable=v1)
    t1.place(relx=0.25,rely=0.1,relwidth=0.2,relheight=0.1)
    lb2 = tk.Label(winNew,text='����/ѡ���ļ���')
    lb2.place(relx=0.1,rely=0.1)
    for i in range (len(v0)):
        rdi = tk.Radiobutton(winNew,text=str(v0[i]),variable=v1,value=v0[i])
        rdi.pack()    
    btClose=tk.Button(winNew,text='�ر�',command=winNew.destroy)
    btClose.place(relx=0.7,rely=0.5)
    btn = tk.Button(winNew,text="OK",command=okfile)
    btn.place(relx=0.6,rely=0.5)
    
def okfile():
##    print(datapath)
    data=pd.read_excel(datapath)
    product=v1.get()
    classify_product(data, int(product))
    

##��������ѯ����
def number():
    winNew = tk.Toplevel(w)
    winNew.geometry('700x350')
    winNew.title('����������')
    t1=tk.Entry(winNew,width=20,textvariable=v2)
    t1.place(relx=0.25,rely=0.1,relwidth=0.2,relheight=0.1)
    lb1 = tk.Label(winNew,text='����/ѡ����Сֵ')
    lb1.place(relx=0.1,rely=0.1)
    t2=tk.Entry(winNew,width=20,textvariable=v3)
    t2.place(relx=0.25,rely=0.2,relwidth=0.2,relheight=0.1)
    lb2 = tk.Label(winNew,text='����/ѡ�����ֵ')
    lb2.place(relx=0.1,rely=0.2)
    scl = tk.Scale(winNew,orient=tk.VERTICAL,length=200,from_=0,to=100,label='��Сֵ',tickinterval=10,resolution=1,variable=v2)
    scl.place(relx=0.1,rely=0.4)
    sc2 = tk.Scale(winNew,orient=tk.VERTICAL,length=200,from_=0,to=200,label='���ֵ',tickinterval=20,resolution=1,variable=v3)
    sc2.place(relx=0.3,rely=0.4)

    btClose=tk.Button(winNew,text='�ر�',command=winNew.destroy)
    btClose.place(relx=0.7,rely=0.5)
    btn = tk.Button(winNew,text="OK",command=oknumber)
    btn.place(relx=0.6,rely=0.5)

def oknumber():
    data=pd.read_excel(datapath)
    min = v2.get()
    max = v3.get()
##    print(min,max)
    classify_num(data, int(min), int(max))

    
    

##��Ԫ����ѯ����
def components():
    winNew = tk.Toplevel(w)
    winNew.geometry('700x350')
    winNew.title('��Ԫ������')
    t1=tk.Entry(winNew,width=20,textvariable=v4)
    t1.place(relx=0.25,rely=0.1,relwidth=0.2,relheight=0.1)
    lb2 = tk.Label(winNew,text='����/ѡ��Ԫ��')
    lb2.place(relx=0.1,rely=0.1)
    for i in range (len(com)):
        rdi = tk.Radiobutton(winNew,text=str(com[i]),variable=v4,value=com[i])
        rdi.pack() 
    btClose=tk.Button(winNew,text='�ر�',command=winNew.destroy)
    btClose.place(relx=0.7,rely=0.5)
    btn = tk.Button(winNew,text="OK",command=okcomponents)
    btn.place(relx=0.6,rely=0.5)
def okcomponents():
    data=pd.read_excel(datapath)
    type = v4.get()
    classify_type(data, type)


##��������ѯ����    
def multiple():
    winNew = tk.Toplevel(w)
    winNew.geometry('1000x600')
    winNew.title('��������ѯ')
    t1=tk.Entry(winNew,width=20,textvariable=v1)
    t1.place(relx=0.25,rely=0.1,relwidth=0.2,relheight=0.1)
    lb2 = tk.Label(winNew,text='����/ѡ���ļ���')
    lb2.place(relx=0.1,rely=0.1)
    for i in range (len(v0)):
        rdi = tk.Radiobutton(winNew,text=str(v0[i]),variable=v1,value=v0[i])
        rdi.place(relx=0.6,rely=0.07*i)
    
    t2=tk.Entry(winNew,width=20,textvariable=v2)
    t2.place(relx=0.25,rely=0.2,relwidth=0.2,relheight=0.1)
    lb2 = tk.Label(winNew,text='����/ѡ����Сֵ')
    lb2.place(relx=0.1,rely=0.2)
    t3=tk.Entry(winNew,width=20,textvariable=v3)
    t3.place(relx=0.25,rely=0.3,relwidth=0.2,relheight=0.1)
    lb3 = tk.Label(winNew,text='����/ѡ�����ֵ')
    lb3.place(relx=0.1,rely=0.3)
    scl = tk.Scale(winNew,orient=tk.VERTICAL,length=200,from_=0,to=100,label='��Сֵ',tickinterval=10,resolution=1,variable=v2)
    scl.place(relx=0.2,rely=0.6)
    sc2 = tk.Scale(winNew,orient=tk.VERTICAL,length=200,from_=0,to=200,label='���ֵ',tickinterval=20,resolution=1,variable=v3)
    sc2.place(relx=0.4,rely=0.6)
    t4=tk.Entry(winNew,width=20,textvariable=v4)
    t4.place(relx=0.25,rely=0.4,relwidth=0.2,relheight=0.1)
    lb4 = tk.Label(winNew,text='����/ѡ��Ԫ��')
    lb4.place(relx=0.1,rely=0.4)
    for i in range (len(com)):
        rdi = tk.Radiobutton(winNew,text=str(com[i]),variable=v4,value=com[i])
        rdi.place(relx=0.8,rely=0.07*i) 
    
    btClose=tk.Button(winNew,text='�ر�',command=winNew.destroy)
    btClose.place(relx=0.8,rely=0.8)
    btn = tk.Button(winNew,text="OK",command=okmultiple)
    btn.place(relx=0.7,rely=0.8)

def okmultiple():
    data=pd.read_excel(datapath)
    product = v1.get()
    min = v2.get()
    max = v3.get()
    type = v4.get()
    classify_multi(data, int(product), int(min), int(max), type)

    
##�����ļ��У����������ļ�·��
def traverse(f):
    #list���ļ���
    list = []
    fs = os.listdir(f)
    for f1 in fs:
        tmp_path = os.path.join(f, f1)
        if not os.path.isdir(tmp_path):
            list.append(tmp_path)
    return list


##ѡ���ļ��в������ݽ��г�������
def folder():
    # ��ѡ���ļ��жԻ���
    # ѡ���ļ���
    Folderpath = filedialog.askdirectory() 
    # ѡ���ļ�
    #Filepath = filedialog.askopenfilename() 
    # ��ӡ�ļ���·��
    #print('Folderpath:', Folderpath)
    x=Folderpath
    l.config(text='��ѡ���·���ǣ�'+str(x))
    
    list = traverse(x);
    
    ##����������ʾ·���е��ļ�
    scroll = tk.Scrollbar()
    text = tk.Text(w, width=45, height=60)
    scroll.pack(side=tk.RIGHT, fill=tk.Y)
    text.pack(side=tk.LEFT, fill=tk.Y)
    scroll.config(command=text.yview)
    text.config(yscrollcommand=scroll.set)
    text.place(relx=0.1,rely=0.1)

    ##��doc�ļ�תΪdocx�ļ�
    for i in list:
        k=len(i)
        word = wc.Dispatch("Word.Application")
        if re.match(r'[^A]*doc',i):
            if re.match(r'[^A]*docx',i):
                k=len(i)
            else:
                doc = word.Documents.Open(i)
                j = i[1+len(x):(k-4)]
                doc.SaveAs(x+'/{}.docx'.format(j), 12)#12Ϊdocx
                doc.Close()
                word.Quit()
    ##�������ļ�תΪ���з�ʽ��ͬ��Excel�ļ�
    #�½����
##    workbook = Workbook()
##    booksheet = workbook.active
##    booksheet.append(["Item",
##                          "Comment",
##                          "Discription",
##                          "Designator",
##                          "Footprint",
##                          "LibRef",
##                          "Quantity",
##                          "Manufacturer",
##                          "Note",
##                          "Level",
##                          "Product"])
##    count = 0

                
    for i in list:#�����ļ���
        
        #j=i[31:35]
        k=len(i)
        
        #print(i)
        test = i
        if not os.path.isdir('test'): #�ж��Ƿ����ļ��У������ļ��вŴ�
            if re.match(r'[^A]*xlsx',test):
                text.insert(tk.INSERT, test+'\n')
                j = i[1+len(x):(k-5)]
                v0.append(j)
##                path=test
##                t = pd.read_excel(path) 
##                for c in range(4, len(t)):
##                    count +=1
##                    list = []
##                    list.append(count)
##                    list.append(t.iloc[c,3] )                 
##                    list.append('')
##                    list.append(t.iloc[c,2] )
##                    list.append(t.iloc[c,4] )
##                    list.append('')
##                    list.append(t.iloc[c,1] )
##                    list.append('')
##                    list.append('' )
##                    list.append('')
##                    list.append(float(j))
##
##                    booksheet.append(list)
   
            if re.match(r'[^A]*txt',test):
                text.insert(tk.INSERT, test+'\n')
                j = i[1+len(x):(k-4)]
                v0.append(j)
##                file = open(test)
##                file_data = pd.read_table(file) #���ļ�
##                file_data.to_csv(x+'/{}.csv'.format(j),index = False)

            if re.match(r'[^A]*docx',test):
                
                text.insert(tk.INSERT, test+'\n')
                j = i[1+len(x):(k-5)]
                v0.append(j)
                path = test
##                f = docx.Document(path)
##                t = f.tables #��ȡ�ļ��еı��
##                table = f.tables[0  ]#��ȡ�ļ��еĵ�һ�����
##                for c in range(1, len(t)):
##                    count +=1
##                    
##                    list = []
##                    list.append(count)
##                    list.append(t.cell(c,2).text)                 
##                    list.append('')
##                    list.append(t.cell(c,0).text)
##                    list.append('')
##                    list.append('')
##                    list.append(t.cell(c,4).text)
##                    list.append('')
##                    list.append(t.cell(c,5).text)
##                    list.append(t.cell(c,3).text)
##                    list.append(float(j))
##
##                    booksheet.append(list)
    if not os.path.exists(x+'/newboms'):
        os.mkdir(x+'/newboms')
##    workbook.save(x+'/newboms/0001.xlsx')
    global datapath
    datapath=x+'/newboms/xxxx.xlsx'
##    print(datapath)
        

##����ͼ�ν���
w=tk.Tk() 
w.geometry('500x300') 
w.title("�������ܷ���ϵͳ") 
l = tk.Label(w, width=50, text='')
l.pack()
m=tk.Menu(w) 
w.config(menu=m) 
insesrtmenu=tk.Menu(m) 
m.add_cascade(label="�����ļ�",menu=insesrtmenu) 
insesrtmenu.add_command(label=" ѡ���ļ���",command=folder)

fmenu=tk.Menu(m) 
m.add_cascade(label="����",menu=fmenu) 
fmenu.add_command(label="���ļ�����",command=file) 
fmenu.add_command(label="����������",command=number) 
fmenu.add_command(label="��Ԫ������",command=components)
fmenu.add_command(label="��������ѯ",command=multiple)

##�������
v0=[]
com=['R','C','L','U','J','T','X']
v1=tk.Variable()
v2=tk.Variable()
v3=tk.Variable()
v4=tk.Variable()
datapath=tk.Variable()

data=tk.Variable()



w.mainloop()
